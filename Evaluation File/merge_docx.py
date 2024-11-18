import os
from docx import Document

def merge_docx_files(file_list, output_file):
    # Create a new Document object for the merged file
    merged_document = Document()

    # Loop through the list of provided files
    for filename in file_list:
        if filename.endswith('.docx'):
            file_path = os.path.join(os.getcwd(), filename)  # Full path to the file
            if os.path.exists(file_path):  # Ensure the file exists
                doc = Document(file_path)

                # Copy each paragraph and its formatting
                for paragraph in doc.paragraphs:
                    # Create a new paragraph in the merged document
                    new_paragraph = merged_document.add_paragraph()

                    # Copy each run (portion of text with the same formatting)
                    for run in paragraph.runs:
                        new_run = new_paragraph.add_run(run.text)

                        # Copy the formatting (font, color, etc.)
                        new_run.font.bold = run.font.bold
                        new_run.font.italic = run.font.italic
                        new_run.font.underline = run.font.underline
                        new_run.font.size = run.font.size
                        new_run.font.name = run.font.name
                        new_run.font.color.rgb = run.font.color.rgb

                # Optionally add a page break between documents
                merged_document.add_page_break()

    # Save the merged document
    merged_document.save(output_file)
    print(f"All files merged into '{output_file}' successfully.")

# Usage
file_list = [
    "Data Mining LAB _ COVER.docx",
    "Data Mining LAB _ Expt 2.docx",
    "Data Mining LAB _ Expt 3.docx",
    "Data Mining LAB _ Expt 4.docx",
    "Data Mining LAB _ Expt 5.docx",
    "Data Mining LAB _ Expt 6.docx",
    "Data Mining LAB _ Expt 7.docx",
    "Data Mining LAB _ Expt 8.docx",
    "Data Mining LAB _ Expt 9.docx",
    "Data Mining LAB _ Expt 10A.docx",
    "Data Mining LAB _ Expt 10B.docx"
]
output_file = './DataMining_12212070.docx'  # Output file name
merge_docx_files(file_list, output_file)